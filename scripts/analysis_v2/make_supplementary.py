#!/usr/bin/env python3
# make_supplementary.py  — builds Supplementary_Tables.docx from frozen files
# needs python-docx:  pip install --user python-docx
import os
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

BASE="/data/habib/metagenome/biosurfactant"
M="."
OUTdoc=f"{BASE}/analysis_v2_community_abundance/Supplementary_Tables.docx"

doc=Document()
for s in doc.sections:
    s.left_margin=s.right_margin=Pt(54)
title=doc.add_paragraph(); r=title.add_run("Supplementary Tables"); r.bold=True; r.font.size=Pt(14)
doc.add_paragraph()

def read_tsv(path, keep_cols=None, max_rows=None):
    if not os.path.exists(path):
        print("MISSING:", path); return None
    rows=[]
    with open(path) as f:
        for i,l in enumerate(f):
            rows.append(l.rstrip("\n").split("\t"))
            if max_rows and i>=max_rows: break
    if keep_cols is not None and rows:
        idx=[rows[0].index(c) for c in keep_cols if c in rows[0]]
        rows=[[r[j] for j in idx] for r in rows]
    return rows

def add_table(num, caption, rows, italic0=False):
    p=doc.add_paragraph()
    rr=p.add_run(f"Table S{num}. "); rr.bold=True; rr.font.size=Pt(9)
    rc=p.add_run(caption); rc.font.size=Pt(9)
    if not rows:
        w=doc.add_paragraph().add_run("[source file not found on disk — see console]"); w.italic=True
        doc.add_paragraph(); return
    t=doc.add_table(rows=len(rows), cols=len(rows[0]))
    try: t.style="Light Grid Accent 1"
    except: t.style="Table Grid"
    for i,row in enumerate(rows):
        for j,val in enumerate(row):
            c=t.cell(i,j); c.text=""
            run=c.paragraphs[0].add_run(str(val)); run.font.size=Pt(8)
            if i==0: run.bold=True
            elif j==0 and italic0: run.italic=True
    doc.add_paragraph()

# ---- S1: sequencing / assembly stats (adjust path/cols to your real file) ----
s1=read_tsv("read_assembly_stats.tsv")
add_table(1,"Sequencing, quality-filtering, and assembly statistics per sample.", s1)

# ---- S2: MAG quality, size, taxonomy ----
add_table(2,"Metagenome-assembled genome (MAG) quality, size, GC content, and GTDB taxonomy.",
          read_tsv(f"genome_economics.tsv"))

# ---- S5: NRPS condensation-domain classification (biosurfactant evidence) ----
s5=read_tsv(f"nrps_classified.tsv", max_rows=None)
if s5 and s5[0][0].startswith(("C3","C6","H2","H6")):   # headerless -> add header
    s5=[["sample","contig","starterC_count","nAMP","bin","season","label"][:len(s5[0])]]+s5
add_table(5,"Classification of NRPS condensation domains. The starter (C-Starter) domain count is diagnostic of lipopeptide assembly.", s5)

# ---- S6: CAZy family abundance by season ----
add_table(6,"Carbohydrate-active enzyme (CAZy) family abundance (whole-community TPM) by season.",
          read_tsv(f"cazy_family_tpm.tsv"))

# ---- S7: verified surface-active-gene carriers (curated) ----
s7=[["Genus","MAG","Function","Gene","KO","PFAM","Status"],
 ["Moraxella_A","H6_concoct_8","EPS/capsule export (Wzy-dependent)","wza, wzb, wzc","K01991, K01104, K16692","Poly_export; LMWPc; Wzz","Verified (contiguous operon)"],
 ["Moraxella_A","H6_concoct_8","Phospholipase A","pldA","K01058","PLA1","Verified"],
 ["Moraxella_A","H6_concoct_8","Fatty-acid desaturase","desC","K00507","FA_desaturase","Verified"],
 ["Loktanella","H6_concoct_40","EPS/capsule export (group-2, ABC)","kpsE, kpsT, exoP","K01992, K09689, K16554","Wzz; ABC_tran","Verified"],
 ["Loktanella","H6_concoct_40","Lysophospholipase","pldB","K01048","Hydrolase_4","Verified"],
 ["Roseovarius","H2_concoct_19","EPS/capsule export (group-2, ABC)","kpsE, kpsT","K01992, K09689","ABC_tran","Verified"],
 ["Roseovarius","H2_concoct_19","Lysophospholipase","pldB","K01048","Hydrolase_4","Verified"]]
add_table(7,"Genome-resolved carriers of surface-active-lipid genes, confirmed at the KO and PFAM level.", s7, italic0=True)

# ---- S8: proteome acidity per MAG ----
add_table(8,"Proteome acidity (genome-wide % acidic minus % basic residues) per MAG; higher values indicate the salt-in strategy.",
          read_tsv(f"proteome_acidity.tsv"))

# ---- note pointing to the large data files ----
p=doc.add_paragraph()
r=p.add_run("Supplementary Data. "); r.bold=True; r.font.size=Pt(9)
r2=p.add_run("Table S3 (whole-community KEGG-ortholog TPM) and Table S4 (all 1,468 biosynthetic gene clusters with taxonomy and coverage) are provided as separate tab-delimited data files owing to their size.")
r2.font.size=Pt(9)

doc.save(OUTdoc)
print("WROTE:", OUTdoc)
